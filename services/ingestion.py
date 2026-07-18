import uuid
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from config.settings import settings
from core.enums import DocumentStatus, OperationType
from core.models import Document
from db.connection import DEFAULT_DB_PATH
from db.repositories import documents_repo
from embeddings.base import EmbeddingProvider
from embeddings.registry import get_embedding_provider
from metadata.logger import log_call
from rag.chunking import chunk
from rag.vector_store import ChromaVectorStore, VectorStore

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


class IngestionError(Exception):
    pass


def ingest_document(
    file_bytes: bytes,
    filename: str,
    title: str,
    topic: str | None = None,
    tags: list[str] | None = None,
    embedding_provider: EmbeddingProvider | None = None,
    embedding_model: str | None = None,
    vector_store: VectorStore | None = None,
    db_path: Path | str = DEFAULT_DB_PATH,
) -> Document:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise IngestionError(
            f"Unsupported file type {extension!r} -- expected one of {sorted(SUPPORTED_EXTENSIONS)}"
        )

    text = _extract_text(file_bytes, extension)
    if not text.strip():
        raise IngestionError(f"No extractable text found in {filename!r}")

    try:
        embedding_provider = embedding_provider or get_embedding_provider()
    except ValueError as exc:
        raise IngestionError(f"Embedding provider unavailable: {exc}") from exc
    embedding_model = embedding_model or settings.default_embedding_model
    vector_store = vector_store or ChromaVectorStore(persist_dir=settings.chroma_persist_dir)

    document = Document(
        id=str(uuid.uuid4()),
        title=title,
        original_filename=filename,
        status=DocumentStatus.INGESTED,
        chunk_count=0,
        topic=topic,
        tags=tags or [],
        created_at=datetime.now(UTC),
    )
    documents_repo.insert(document, db_path=db_path)

    try:
        chunks = chunk(text, document_id=document.id)
        documents_repo.update_status(document.id, DocumentStatus.CHUNKED, db_path=db_path)

        result = embedding_provider.embed([c.text for c in chunks], model=embedding_model)
        log_call(
            operation_type=OperationType.EMBEDDING,
            provider=result.provider,
            model=result.model,
            input_tokens=result.input_tokens,
            output_tokens=0,
            latency_ms=result.latency_ms,
            cost_usd=result.cost_usd,
            db_path=db_path,
        )
        documents_repo.update_status(document.id, DocumentStatus.EMBEDDED, db_path=db_path)

        vector_store.add(
            ids=[f"{document.id}_{c.chunk_index}" for c in chunks],
            vectors=result.vectors,
            metadata=[
                {
                    "document_id": c.document_id,
                    "chunk_index": c.chunk_index,
                    "text": c.text,
                    "start_char": c.start_char,
                    "end_char": c.end_char,
                }
                for c in chunks
            ],
        )
        documents_repo.update_status(
            document.id, DocumentStatus.READY, chunk_count=len(chunks), db_path=db_path
        )
    except Exception as exc:
        documents_repo.update_status(document.id, DocumentStatus.FAILED, db_path=db_path)
        raise IngestionError(f"Ingestion failed for document {document.id}: {exc}") from exc

    return documents_repo.get(document.id, db_path=db_path)


def _extract_text(file_bytes: bytes, extension: str) -> str:
    try:
        if extension == ".pdf":
            reader = PdfReader(BytesIO(file_bytes))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        docx_document = DocxDocument(BytesIO(file_bytes))
        return "\n\n".join(p.text for p in docx_document.paragraphs)
    except Exception as exc:
        raise IngestionError(f"Failed to extract text from file: {exc}") from exc
